import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ACCENT = RGBColor(14, 74, 122)
DARK = RGBColor(34, 34, 34)
MUTED = RGBColor(104, 113, 122)


PROFILE = {
    "name": "Shuo Sun",
    "title": "Postdoctoral Associate",
    "affiliations": "SMART M3S IRG · MIT Center for Collective Intelligence",
    "email": "shuo.sun@smart.mit.edu",
    "github": "github.com/SS47816",
    "linkedin": "linkedin.com/in/shuosunss",
    "scholar": "scholar.google.com/citations?user=xe6YcxYAAAAJ",
    "summary": (
        "Research at the intersection of robotics, autonomous systems, and human-AI "
        "collective intelligence, with an emphasis on systems that augment human "
        "capability and support the future of work and education."
    ),
    "research_interests": [
        "Robotics and Autonomous Systems",
        "Human-AI Collective Intelligence",
        "Motion Planning and Prediction",
        "Scenario Generation and Simulation",
    ],
}


APPOINTMENTS = [
    {
        "period": "Jun 2025 - Present",
        "title": "Postdoctoral Associate",
        "institution": "Singapore-MIT Alliance for Research and Technology (SMART)",
        "detail": "M3S Interdisciplinary Research Group · MIT Center for Collective Intelligence",
        "bullets": [
            "Research at the intersection of robotics and human-AI collective intelligence.",
            "Developing intelligent systems that augment human capability, support collective intelligence, and empower the future of work and education.",
            "Principal Investigator: Professor Thomas W. Malone (MIT Sloan School of Management).",
        ],
    },
    {
        "period": "Aug 2020 - May 2025",
        "title": "Graduate Tutor & Project Engineer",
        "institution": "National University of Singapore",
        "detail": "Mechanical Engineering · NUS Advanced Robotics Centre",
        "bullets": [
            "Led design and deployment across the Autonomous Road Sweeper and Autonomous Mini-bus projects.",
            "Designed a full-stack autonomous driving software system for semi-structured environments.",
            "Developed course content for ME5413 Robotics, including hands-on homework, project challenges, tutorials, and consultation sessions.",
        ],
    },
]


EDUCATION = [
    {
        "period": "2021 - 2025",
        "title": "Doctor of Philosophy (Ph.D.) in Mechanical Engineering",
        "institution": "National University of Singapore",
        "detail": "Supervisor: Professor Marcelo H. Ang Jr. · Research: autonomous systems, motion prediction, motion planning, and simulation.",
    },
    {
        "period": "2016 - 2020",
        "title": "Bachelor of Engineering (Honours) in Mechanical Engineering",
        "institution": "National University of Singapore",
        "detail": "Double major in Mechanical Engineering and the Innovation & Design Programme (iDP). Science & Technology Undergraduate Scholarship (2016 - 2020).",
    },
]

PUBLICATIONS = [
    '[1] A. Cai*, I. Yeckehzaare*, S. Sun*, V. Charisi*, X. Wang, A. Imran, R. Laubacher, A. Prakash, and T. Malone, "Where can AI be used? Insights from a deep ontology of work activities," arXiv, 2026.',
    '[2] S. Sun, Y. Zhao, C. D. W. Lee, J. Sun, C. Yuan, Z. Huang, D. Li, J. K. Yeoh, A. Prakash, T. W. Malone et al., "AGI-Elo: How far are we from mastering a task?" Advances in Neural Information Processing Systems (NeurIPS 2025), 2025.',
    '[3] S. Sun, Z. Gu, T. Sun, J. Sun, C. Yuan, Y. Han, D. Li, and M. H. Ang, "DriveSceneGen: Generating diverse and realistic driving scenarios from scratch," IEEE Robotics and Automation Letters (RA-L 2024 & ICRA 2025), 2024.',
    '[4] S. Sun, J. Chen, J. Sun, C. Yuan, Y. Li, T. Zhang, and M. H. Ang, "FISS+: Efficient and focused trajectory generation and refinement using fast iterative search and sampling strategy," IROS 2023.',
    '[5] S. Sun, Z. Liu, H. Yin, and M. H. Ang, "FISS: A trajectory planning framework using fast iterative search and sampling strategy for autonomous driving," IEEE Robotics and Automation Letters (RA-L 2022), 2022.',
    '[6] H. Yang, S. Ding, J. Wang, S. Sun, R. Swaminathan, S. W. L. Ng, X. Pan, and G. W. Ho, "Computational design of ultra-robust strain sensors for soft robot perception and autonomy," Nature Communications, 2024.',
]


AWARDS = [
    "2nd place, 2025 Waymo Open Dataset Interaction Prediction Challenge.",
    "3rd place, 2024 Waymo Open Dataset Motion Prediction Challenge.",
]


TEACHING = [
    "Developed course content for ME5413 Robotics, including annual homework and final-project challenge design, tutorials, and consultation sessions.",
    "Supervised 30+ Master's and 40+ Bachelor's theses and research projects.",
]


SKILLS = [
    ("Robotics", "autonomous driving, object detection, sensor fusion, motion prediction, motion planning, motion control, scenario generation, and simulation"),
    ("AI/ML", "CNNs, transformers, diffusion models, LLMs, multi-agent systems, and Graph-RAG"),
    ("Programming", "C/C++, Python, PyTorch, TensorFlow, ROS, OpenCV, PCL, Qt, and MATLAB"),
]


def set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tbl_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tbl_borders.append(element)
        element.set(qn("w:val"), "nil")


def add_bottom_border(paragraph, color="C9D4E3", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.first_child_found_in("w:pBdr")
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0E4A7A")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    r_pr.append(underline)
    new_run.append(r_pr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_run_font(run, font_name, size, bold=False, italic=False, color=DARK):
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_section_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title.upper())
    set_run_font(run, "Arial", 10.5, bold=True, color=ACCENT)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_bottom_border(p)


def add_intro(doc):
    name = doc.add_paragraph()
    name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name.paragraph_format.space_after = Pt(2)
    run = name.add_run(PROFILE["name"])
    set_run_font(run, "Arial", 24, bold=True)

    role = doc.add_paragraph()
    role.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    role.paragraph_format.space_after = Pt(2)
    role_run = role.add_run(f'{PROFILE["title"]} · {PROFILE["affiliations"]}')
    set_run_font(role_run, "Arial", 10.5, color=ACCENT)

    contact = doc.add_paragraph()
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact.paragraph_format.space_after = Pt(8)
    add_hyperlink(contact, PROFILE["email"], f'mailto:{PROFILE["email"]}')
    for label, url in (
        ("Google Scholar", f'https://{PROFILE["scholar"]}'),
        ("GitHub", f'https://{PROFILE["github"]}'),
        ("LinkedIn", f'https://{PROFILE["linkedin"]}'),
    ):
        run = contact.add_run(" | ")
        set_run_font(run, "Arial", 9.5, color=MUTED)
        add_hyperlink(contact, label, url)

    summary = doc.add_paragraph()
    summary.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    summary.paragraph_format.space_after = Pt(8)
    summary.paragraph_format.left_indent = Inches(0.35)
    summary.paragraph_format.right_indent = Inches(0.35)
    summary_run = summary.add_run(PROFILE["summary"])
    set_run_font(summary_run, "Georgia", 10.5, italic=True, color=MUTED)

    interests = doc.add_paragraph()
    interests.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    interests.paragraph_format.space_after = Pt(8)
    label = interests.add_run("Research Interests: ")
    set_run_font(label, "Arial", 9.5, bold=True, color=ACCENT)
    values = interests.add_run(" · ".join(PROFILE["research_interests"]))
    set_run_font(values, "Georgia", 10.5)


def add_entry_table(doc, item, include_bullets=True):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    remove_table_borders(table)

    left_cell, right_cell = table.rows[0].cells
    set_cell_width(left_cell, Inches(1.45))
    set_cell_width(right_cell, Inches(5.55))
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    left_p = left_cell.paragraphs[0]
    left_p.paragraph_format.space_after = Pt(0)
    left_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    left_run = left_p.add_run(item["period"])
    set_run_font(left_run, "Arial", 9.5, color=MUTED)

    right_p = right_cell.paragraphs[0]
    right_p.paragraph_format.space_after = Pt(0)
    title_run = right_p.add_run(item["title"])
    set_run_font(title_run, "Georgia", 11.5, bold=True)
    title_run.add_break(WD_BREAK.LINE)
    inst_run = right_p.add_run(item["institution"])
    set_run_font(inst_run, "Arial", 10, color=ACCENT)

    if item.get("detail"):
        detail_p = right_cell.add_paragraph()
        detail_p.paragraph_format.space_after = Pt(2)
        detail_run = detail_p.add_run(item["detail"])
        set_run_font(detail_run, "Georgia", 10.2, italic=True, color=MUTED)

    if include_bullets:
        for bullet in item.get("bullets", []):
            p = right_cell.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(bullet)
            set_run_font(run, "Georgia", 10.2)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_publications(doc):
    for citation in PUBLICATIONS:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(citation)
        set_run_font(run, "Georgia", 10.2)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(2)
    note_run = note.add_run("* Equal contribution.")
    set_run_font(note_run, "Georgia", 9.5, italic=True, color=MUTED)


def add_simple_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(item)
        set_run_font(run, "Georgia", 10.2)


def add_skills(doc):
    for label, value in SKILLS:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, "Arial", 10, bold=True, color=ACCENT)
        value_run = p.add_run(value)
        set_run_font(value_run, "Georgia", 10.2)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(16)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.start_type = WD_SECTION_START.NEW_PAGE

    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style_r_pr = style._element.get_or_add_rPr()
    style_r_fonts = style_r_pr.rFonts
    if style_r_fonts is None:
        style_r_fonts = OxmlElement("w:rFonts")
        style_r_pr.append(style_r_fonts)
    style_r_fonts.set(qn("w:ascii"), "Georgia")
    style_r_fonts.set(qn("w:hAnsi"), "Georgia")
    style_r_fonts.set(qn("w:eastAsia"), "Georgia")
    style.font.size = Pt(10.5)


def load_selected_projects(root):
    project_data = json.loads((root / "data" / "projects.json").read_text())
    projects_by_id = {item["id"]: item for item in project_data.get("items", [])}
    selected_projects = []

    for project_id in project_data.get("collections", {}).get("academicCv", []):
        project = projects_by_id.get(project_id)
        if not project:
            continue
        selected_projects.append(
            {
                "period": project.get("period", ""),
                "title": project.get("title", ""),
                "institution": project.get("org", ""),
                "bullets": project.get("contributions", []),
            }
        )

    return selected_projects


def build_document(projects):
    doc = Document()
    configure_document(doc)

    add_intro(doc)

    add_section_heading(doc, "Appointments")
    for item in APPOINTMENTS:
        add_entry_table(doc, item)

    add_section_heading(doc, "Education")
    for item in EDUCATION:
        add_entry_table(doc, item, include_bullets=False)

    add_section_heading(doc, "Selected Projects")
    for item in projects:
        add_entry_table(doc, item)

    add_section_heading(doc, "Selected Publications")
    add_publications(doc)

    add_section_heading(doc, "Awards and Distinctions")
    add_simple_bullets(doc, AWARDS)

    add_section_heading(doc, "Teaching and Mentoring")
    add_simple_bullets(doc, TEACHING)

    add_section_heading(doc, "Technical Skills")
    add_skills(doc)

    return doc


def main():
    root = Path(__file__).resolve().parents[1]
    output = root / "files" / "Shuo_SUN_CV_academic.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    projects = load_selected_projects(root)
    doc = build_document(projects)
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
